"""Build the merge-field-only BreachBrief SLA credit memo template.

The agent resolves and flattens the variable-length incident audit. Doctavian
shapes the resulting facts into a real document using the direct-field surface
proven against the sponsor demo tenant.
"""
from __future__ import annotations

import pathlib
import re
from zipfile import ZipFile, ZipInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = pathlib.Path(__file__).resolve().parent / "sla-credit-memo.docx"
PACKAGE_TIMESTAMP = (1980, 1, 1, 0, 0, 0)


def para(doc, text: str, *, bold: bool = False, size: int = 10, space_after: int = 4):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def normalize_package(path: pathlib.Path) -> None:
    """Remove build-time metadata and make the DOCX byte-reproducible."""
    with ZipFile(path) as source:
        members = []
        for item in source.infolist():
            name = item.filename
            if name.startswith(("docProps/", "customXml/")):
                continue
            payload = source.read(name)
            if name == "_rels/.rels":
                payload = re.sub(
                    br'<Relationship[^>]+Target="docProps/[^\"]+"[^>]*/>',
                    b"",
                    payload,
                )
            elif name == "word/_rels/document.xml.rels":
                payload = re.sub(
                    br'<Relationship[^>]+Target="\.\./customXml/[^\"]+"[^>]*/>',
                    b"",
                    payload,
                )
            elif name == "[Content_Types].xml":
                payload = re.sub(
                    br'<Override[^>]+PartName="/(?:docProps|customXml)/[^\"]+"[^>]*/>',
                    b"",
                    payload,
                )
                payload = re.sub(
                    br'<Default[^>]+Extension="jpeg"[^>]*/>', b"", payload
                )
            if name.endswith((".xml", ".rels")):
                payload = re.sub(br'\s+w:rsid\w*="[^\"]*"', b"", payload)
                payload = re.sub(br'<w:rsids>.*?</w:rsids>', b"", payload, flags=re.DOTALL)
                payload = re.sub(br'<w:rsid\s+[^>]*/>', b"", payload)
                payload = re.sub(br'<w:savePreviewPicture\s*/>', b"", payload)
                payload = re.sub(br'<w1[45]:docId[^>]*/>', b"", payload)
                payload = re.sub(
                    br'\s+w1[45]:(?:paraId|textId)="[^\"]*"', b"", payload
                )
            members.append((name, payload, item.compress_type))
    temporary = path.with_name("." + path.name + ".tmp")
    with ZipFile(temporary, "w") as target:
        for name, payload, compression in members:
            item = ZipInfo(name, date_time=PACKAGE_TIMESTAMP)
            item.compress_type = compression
            item.create_system = 3
            mode = 0o700 if name.endswith("/") else 0o600
            item.external_attr = mode << 16
            target.writestr(item, payload, compresslevel=9)
    temporary.replace(path)
    path.chmod(0o600)


def build(out: pathlib.Path = OUT) -> pathlib.Path:
    doc = Document()
    doc.core_properties.author = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.subject = ""
    doc.core_properties.title = ""
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_heading("Service Level Credit Memo", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para(doc, "{!Account[0].Name} — {!Account[0].Tier} agreement", bold=True, size=12)
    para(doc,
         "Billing period {!Account[0].PeriodStart} to {!Account[0].PeriodEnd}  ·  "
         "contract availability {!Account[0].ContractUptimePct}%  ·  "
         "monthly fee ${!Account[0].MonthlyFeeUsd}")

    doc.add_heading("Reconciled outcome", level=1)
    para(doc, "{!Account[0].Outcome}", bold=True)
    para(doc,
         "Services reviewed: {!Account[0].ServiceCount}  ·  "
         "services in breach: {!Account[0].BreachCount}  ·  "
         "excluded source rows: {!Account[0].ExcludedCount}")
    para(doc, "Regulatory status: {!Account[0].RegulatoryStatus}")

    doc.add_heading("Availability by service", level=1)
    para(doc, "{!Account[0].ServiceAudit}")

    doc.add_heading("Reconciled outage audit", level=1)
    para(doc, "{!Account[0].OutageAudit}")

    doc.add_heading("Excluded source rows", level=1)
    para(doc, "{!Account[0].ExclusionAudit}")

    doc.add_heading("Total", level=1)
    para(doc,
         "Period minutes: {!Account[0].PeriodMinutes}  ·  "
         "reconciled downtime: {!Account[0].TotalDowntimeMinutes} minutes")
    para(doc, "Credit payable: ${!Account[0].TotalCreditUsd}", bold=True, size=12)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    normalize_package(out)
    return out


if __name__ == "__main__":
    path = build()
    print(f"template written: {path} ({path.stat().st_size} bytes)")
